- 👋 Hi, I’m @iMath247
- 👀 I’m interested in ...
- 🌱 I’m currently learning ...
- 💞️ I’m looking to collaborate on ...
- 📫 How to reach me ...
- 😄 Pronouns: ...
- ⚡ Fun fact: ...

<!---
iMath247/iMath247 is a ✨ special ✨ repository because its `README.md` (this file) appears on your GitHub profile.
You can click the Preview link to take a look at your changes.
--->
import datetime
from docx import Document

# Create a new Document object
doc = Document(r'C:\VTC\Dropbox\PYTHON\Template-import\templateL45.docx')

# Check if the first paragraph is empty and delete it
if not doc.paragraphs[0].text:
    p = doc.paragraphs[0]._element
    p.getparent().remove(p)
    # Remove the assignment to _p and _element
# Add a title using a defined style in the template
#

###########################################################
doc.add_heading("BÀI 2. CÁC PHÉP TÍNH VỀ PHÂN SỐ", level=0)

doc.add_heading("NĂM HỌC 2024 - 2025", level=6)

# NỘI DUNG VĂN BẢN
#
doc.add_heading("QUY TẮC", level=1)

#
heading = doc.add_heading("Tổng hai phân số: ", level=4) 
heading.add_run("$\\frac{A}{M} + \\frac{B}{M} = \\frac{A+B}{M}$")

#
heading = doc.add_heading("Hiệu hai phân số: ", level=4)
heading.add_run("$\\frac{A}{M} - \\frac{B}{M} = \\frac{A-B}{M}$")
#
doc.add_paragraph("Nhận xét: Để cộng hoặc trừ hai phân số, ta phải đưa chúng về cùng một mẫu số rồi thực hiện phép cộng hoặc trừ trên tử số.")
# Ví dụ về phép cộng và trừ 2 phân số cùng mẫu ngẫu nhiên
import random
import math
a = random.randint(10, 30)*3
b = random.randint(10, 40)*3
m = random.randint(10, 50)*3
n = random.randint(10, 50)*3
# Thêm điều kiện để kết quả phép toán Tổng hiệu 2 phân số không âm
while a - b <= 0:
    a = random.randint(10, 30)*3
    b = random.randint(10, 40)*3
    m = random.randint(10, 50)*3
    n = random.randint(10, 50)*3
    break

doc.add_paragraph(f"Ví dụ 1: Tính tổng và hiệu của hai phân số sau: $\\frac{{{a}}}{{{m}}} + \\frac{{{b}}}{{{m}}}$ và $\\frac{{{a}}}{{{n}}} - \\frac{{{b}}}{{{n}}}$")
doc.add_paragraph(f"Giải: $\\frac{{{a}}}{{{m}}} + \\frac{{{b}}}{{{m}}}$ = $\\frac{{{a+b}}}{{{m}}}$ và $\\frac{{{a}}}{{{n}}} - \\frac{{{b}}}{{{n}}}$ = $\\frac{{{a-b}}}{{{n}}}$")
doc.add_paragraph("Kết quả: $\\frac{a}{m} + \\frac{b}{m} = \\frac{a+b}{m}$ và $\\frac{a}{n} - \\frac{b}{n} = \\frac{a-b}{n}$")
# Rút gọn kết quả Tổng và Hiệu 2 phân số kết quả phân số tối giản
doc.add_paragraph(f"Kết quả tối giản: $\\frac{{{a}}}{{{m}}} + \\frac{{{b}}}{{{m}}}$ = $\\frac{{{a+b}}}{{{m}}}$ = $\\frac{{{(a+b)//math.gcd(a+b,m)}}}{{{m//math.gcd(a+b,m)}}}$ và $\\frac{{{a}}}{{{n}}} - \\frac{{{b}}}{{{n}}}$ = $\\frac{{{a-b}}}{{{n}}}$ = $\\frac{{{(a-b)//math.gcd(a-b,n)}}}{{{n//math.gcd(a-b,n)}}}$")
# Nhận xét kết quả phân số cần rút gọn tối giản
doc.add_paragraph("Nhận xét: Kết quả Tổng và Hiệu 2 phân số cần rút gọn tối giản")

#
heading = doc.add_heading("Tích hai phân số: ", level=4)
heading.add_run("$\\frac{A}{M} \\times \\frac{B}{N} = \\frac{A \\times B}{M \\times N}$")

heading = doc.add_heading("Thương hai phân số: ", level=4)
heading.add_run("$\\frac{A}{M} : \\frac{B}{N} = \\frac{A}{M} \\times \\frac{N}{B}$")
# Nhận xét phép chia 2 phân số
doc.add_paragraph("Nhận xét: Để chia hai phân số, ta phải đưa phân số chia về dạng đảo ngược TỬ SỐ & MẪU SỐ ('nghịch đảo') rồi nhân với phân số bị chia.")
# Ví dụ 2 phân số ngẫu nhiên về phép toán nhân và chia
a = random.randint(3, 19)*3
b = random.randint(5, 19)*3
m = random.randint(11, 39)*3
n = random.randint(11, 39)*3

doc.add_paragraph(f"Ví dụ 2: Tính tích và thương của hai phân số sau: $\\frac{{{a}}}{{{m}}} \\times \\frac{{{b}}}{{{n}}}$ và $\\frac{{{a}}}{{{m}}} : \\frac{{{b}}}{{{n}}}$")
doc.add_paragraph(f"Giải: $\\frac{{{a}}}{{{m}}} \\times \\frac{{{b}}}{{{n}}}$ = $\\frac{{{a*b}}}{{{m*n}}}$ và $\\frac{{{a}}}{{{m}}} : \\frac{{{b}}}{{{n}}}$ = $\\frac{{{a}}}{{{m}}} \\times \\frac{{{n}}}{{{b}}}$")
doc.add_paragraph("Kết quả: $\\frac{a}{m} \\times \\frac{b}{n} = \\frac{a \\times b}{m \\times n}$ và $\\frac{a}{m} : \\frac{b}{n} = \\frac{a}{m} \\times \\frac{n}{b}$")

# Rút gọn kết quả Tích và Thương 2 phân số kết quả phân số tối giản
doc.add_paragraph(f"Kết quả tối giản: $\\frac{{{a}}}{{{m}}} \\times \\frac{{{b}}}{{{n}}}$ = $\\frac{{{a*b}}}{{{m*n}}}$ = $\\frac{{{(a*b)//math.gcd(a*b,m*n)}}}{{{(m*n)//math.gcd(a*b,m*n)}}}$ và $\\frac{{{a}}}{{{m}}} : \\frac{{{b}}}{{{n}}}$ = $\\frac{{{a}}}{{{m}}} \\times \\frac{{{n}}}{{{b}}}$ = $\\frac{{{a*n}}}{{{m*b}}}$ = $\\frac{{{(a*n)//math.gcd(a*n,m*b)}}}{{{(m*b)//math.gcd(a*n,m*b)}}}$")
doc.add_paragraph("Nhận xét: Kết quả Tích và Thương 2 phân số cần rút gọn chia từng phần trước khi nhân.")



#################################################
doc.add_heading("BÀI TẬP", level=1)
# Câu 1: Tính tổng, hiệu, tích, thương của các phân số sau.
doc.add_heading("Tính tổng, hiệu, tích, thương của các phân số sau.", level=3)

# Function to generate two random fractions and perform addition, subtraction, multiplication, and division
def generate_and_solve_fractions():
    # Generate random numbers for fractions ensuring the result is positive
    while True:
        a = random.randint(10, 30) * 3
        b = random.randint(10, 40) * 3
        m = random.randint(10, 50) * 3
        n = random.randint(10, 50) * 3
        if a - b > 0:
            break

    # Addition
    sum_numerator = a + b
    sum_denominator = m
    simplified_sum_numerator = sum_numerator // math.gcd(sum_numerator, sum_denominator)
    simplified_sum_denominator = sum_denominator // math.gcd(sum_numerator, sum_denominator)
    doc.add_paragraph(f"$M = \\frac{{{a}}}{{{m}}} + \\frac{{{b}}}{{{m}}}$.\nĐáp án: $\\frac{{{simplified_sum_numerator}}}{{{simplified_sum_denominator}}}$")

    # Subtraction
    difference_numerator = a - b
    difference_denominator = n
    simplified_difference_numerator = difference_numerator // math.gcd(difference_numerator, difference_denominator)
    simplified_difference_denominator = difference_denominator // math.gcd(difference_numerator, difference_denominator)
    doc.add_paragraph(f"$N = \\frac{{{a}}}{{{n}}} - \\frac{{{b}}}{{{n}}}$.\nĐáp án: $\\frac{{{simplified_difference_numerator}}}{{{simplified_difference_denominator}}}$")

    # Multiplication
    product_numerator = a * b
    product_denominator = m * n
    simplified_product_numerator = product_numerator // math.gcd(product_numerator, product_denominator)
    simplified_product_denominator = product_denominator // math.gcd(product_numerator, product_denominator)
    doc.add_paragraph(f"$P = \\frac{{{a}}}{{{m}}} \\times \\frac{{{b}}}{{{n}}}$.\nĐáp án: $\\frac{{{simplified_product_numerator}}}{{{simplified_product_denominator}}}$")

    # Division
    quotient_numerator = a * n
    quotient_denominator = m * b
    simplified_quotient_numerator = quotient_numerator // math.gcd(quotient_numerator, quotient_denominator)
    simplified_quotient_denominator = quotient_denominator // math.gcd(quotient_numerator, quotient_denominator)
    doc.add_paragraph(f"$Q = \\frac{{{a}}}{{{m}}} : \\frac{{{b}}}{{{n}}}$.\nĐáp án: $\\frac{{{simplified_quotient_numerator}}}{{{simplified_quotient_denominator}}}$")

# Call the function to generate and solve fractions
generate_and_solve_fractions()
#

#
#
#################################################
# Bài 2. Tính hợp lí
doc.add_heading("Tính hợp lí.", level=3)

# Function to generate a random fraction and simplify it
a=random.randint(10, 30)
b=random.randint(10, 40)
c=random.randint(10, 50)
m=random.randint(2, 9)
n=random.randint(2, 9)
# Thêm điều kiện để kết quả phép toán Tổng hiệu 2 phân số không âm
while a - b <= 0 or m == n or a == b:
    a = random.randint(10, 30)
    b = random.randint(10, 40)
    c = random.randint(10, 50)
    m = random.randint(2, 9)
    n = random.randint(2, 9)
    break
# Tổng 2 phân số a*m/c*m + b*n/c*n
doc.add_paragraph(f"$A = \\frac{{{(a*m)}}}{{{(c*m)}}} + \\frac{{{(b*n)}}}{{{(c*n)}}}$")
# Hiệu 2 phân số a*m/c*m - b*n/c*n
doc.add_paragraph(f"$B = \\frac{{{a*m}}}{{{c*m}}} - \\frac{{{b*n}}}{{{c*n}}}$")
# Đáp án: A và B tối giản
import math  # Đảm bảo rằng thư viện math đã được import

# Giả sử các biến a, b, c, m, n đã được khai báo và có giá trị

# Tính GCD cho kết quả của A và B
gcd_A = math.gcd(a+b, c)
gcd_B = math.gcd(a-b, c)

# Tối giản phân số cho A và B
simplified_A_numerator = (a+b) // gcd_A
simplified_A_denominator = c // gcd_A
simplified_B_numerator = (a-b) // gcd_B
simplified_B_denominator = c // gcd_B

# Giả sử biến doc đã được khởi tạo đúng cách
doc.add_paragraph(f"Đáp án: $A=\\frac{{{simplified_A_numerator}}}{{{simplified_A_denominator}}}$ và $B=\\frac{{{simplified_B_numerator}}}{{{simplified_B_denominator}}}$")


#
import random

def generate_random_1abcd():
    number_str = '1' + ''.join(str(random.randint(0, 1)) for _ in range(4))
    return int(number_str, 10)

def generate_random_1abcde():
    number_str = '1' + ''.join(str(random.randint(0, 1)) for _ in range(5))
    return int(number_str, 10)

m = generate_random_1abcd()
n = generate_random_1abcde()

# Continue with the rest of your code
# Function to generate a random fraction and simplify it
a = random.randint(2, 9)
b = random.randint(2, 9)
c = random.randint(2, 9)

# Thêm điều kiện để kết quả phép toán Tổng hiệu 2 phân số không âm
while a - b < 0 or a == b or b == c:
    a = random.randint(2, 9)
    b = random.randint(2, 9)
    c = random.randint(2, 9)

# Tổng 2 phân số a*m/c*m + b*n/c*n
doc.add_paragraph(f"$C_{{1}} = \\frac{{{int(a*m)}}}{{{int(c*m)}}} + \\frac{{{int(b*n)}}}{{{int(c*n)}}}$")
# Hiệu 2 phân số a*m/c*m - b*n/c*n
m = generate_random_1abcd()
n = generate_random_1abcde()
doc.add_paragraph(f"$C_{{2}} = \\frac{{{a*m}}}{{{c*m}}} - \\frac{{{b*n}}}{{{c*n}}}$")
# Đáp án: A và B tối giản
import math  # Đảm bảo rằng thư viện math đã được import

# Giả sử các biến a, b, c, m, n đã được khai báo và có giá trị

# Tính GCD cho kết quả của A và B
gcd_A = math.gcd(a+b, c)
gcd_B = math.gcd(a-b, c)

# Tối giản phân số cho A và B
simplified_A_numerator = (a+b) // gcd_A
simplified_A_denominator = c // gcd_A
simplified_B_numerator = (a-b) // gcd_B
simplified_B_denominator = c // gcd_B

# Giả sử biến doc đã được khởi tạo đúng cách
doc.add_paragraph(f"Đáp án: $C_{{1}}=\\frac{{{simplified_A_numerator}}}{{{simplified_A_denominator}}}$ và $C_{{2}}=\\frac{{{simplified_B_numerator}}}{{{simplified_B_denominator}}}$")

# Tích 2 phân số a*m/c*m * b*n/c*n
m = generate_random_1abcd()
n = generate_random_1abcde()
doc.add_paragraph(f"$D_{{1}} = \\frac{{{a*m}}}{{{c*m}}} \\times \\frac{{{b*n}}}{{{c*n}}}$")
# Thương 2 phân số a*m/c*m : b*n/c*n
m = generate_random_1abcd()
n = generate_random_1abcde()
doc.add_paragraph(f"$D_{{2}} = \\frac{{{a*m}}}{{{c*m}}}:\\frac{{{b*n}}}{{{c*n}}}$")
# Đáp án tối giản cho D1 và D2
# Tính GCD cho kết quả của D1 và D2
gcd_D1 = math.gcd(a*b, c*c)
gcd_D2 = math.gcd(a*c, b*c)
# Tối giản phân số cho D1 và D2
simplified_D1_numerator = (a*b) // gcd_D1
simplified_D1_denominator = (c*c) // gcd_D1
simplified_D2_numerator = (a*c) // gcd_D2
simplified_D2_denominator = (b*c) // gcd_D2
# In kết quả tối giản
doc.add_paragraph(f"Đáp án: $D_{{1}}=\\frac{{{simplified_D1_numerator}}}{{{simplified_D1_denominator}}}$ và $D_{{2}}=\\frac{{{simplified_D2_numerator}}}{{{simplified_D2_denominator}}}$")


#
import random
from fractions import Fraction

def generate_non_integer_fraction():
    """Tạo một phân số ngẫu nhiên không nguyên."""
    numerator = random.randint(1, 10)
    denominator = random.randint(1, 10)
    while numerator % denominator == 0:  # Đảm bảo phân số không nguyên
        numerator = random.randint(1, 10)
    return Fraction(numerator, denominator)

def fraction_to_latex(fraction):
    """Chuyển đổi phân số sang chuỗi LaTeX."""
    if fraction.denominator == 1:
        return str(fraction.numerator)
    return f"\\frac{{{fraction.numerator}}}{{{fraction.denominator}}}"


# In ra 4 phân số
def calculate_and_print_latex_expression(a, b, c, d):
    # Tìm biểu thức (b operations1 c operations2 d) > 0
    while True:
        operations = [random.choice(['+', '-']) for _ in range(2)]
        operations1, operations2 = operations
        
        # Sử dụng eval để tính toán kết quả tạm thời
        intermediate_result = eval(f"b {operations1} c {operations2} d")
        if intermediate_result > 0:
            break

    # Tính toán biểu thức a*(b operations1 c operations2 d)
    result = eval(f"a * (b {operations1} c {operations2} d)")

    # In ra biểu thức và kết quả dưới dạng LaTeX
    latex_a = fraction_to_latex(a)
    latex_expression = f"{latex_a} \\times {fraction_to_latex(b)} {operations1} {latex_a} \\times  {fraction_to_latex(c)} {operations2} {latex_a} \\times {fraction_to_latex(d)}"
    latex_result = fraction_to_latex(result)
    return latex_expression, latex_result

# Tạo 4 phân số ngẫu nhiên khác nhau và không nguyên
fractions = [generate_non_integer_fraction() for _ in range(4)]
a, b, c, d = fractions
latex_expression, latex_result = calculate_and_print_latex_expression(a, b, c, d)  
doc.add_paragraph(f"$E= {latex_expression}$")
doc.add_paragraph(f"Đáp án: $E= {latex_result}$")
#
def calculate_and_print_latex_expression_1(a, b, c, d):
    # Tìm biểu thức (b operations1 c operations2 d) > 0
    while True:
        operations = [random.choice(['+', '-']) for _ in range(2)]
        operations1, operations2 = operations
        
        # Sử dụng eval để tính toán kết quả tạm thời
        intermediate_result = eval(f"b {operations1} c {operations2} d")
        if intermediate_result > 0:
            break
        break    
    # Tính toán biểu thức a*(b operations1 c operations2 d)
    result = eval(f"a * (b {operations1} c {operations2} d)")

    # In ra biểu thức và kết quả dưới dạng LaTeX
    latex_a = fraction_to_latex(a)
    latex_expression = f" {fraction_to_latex(b)} \\times {latex_a} {operations1}   {fraction_to_latex(c)} \\times {latex_a} {operations2} {fraction_to_latex(d)} \\times {latex_a}"
    latex_result = fraction_to_latex(result)
    return latex_expression, latex_result
# Tạo 4 phân số ngẫu nhiên khác nhau và không nguyên
fractions = [generate_non_integer_fraction() for _ in range(4)]
a, b, c, d = fractions
latex_expression, latex_result = calculate_and_print_latex_expression_1(a, b, c, d)  
doc.add_paragraph(f"$F= {latex_expression}$")
doc.add_paragraph(f"Đáp án: $F= {latex_result}$")
#
def calculate_and_print_latex_expression_2(a, b, c, d):
    # Tìm biểu thức (b operations1 c operations2 d) > 0
    while True:
        operations = [random.choice(['+', '-']) for _ in range(2)]
        operations1, operations2 = operations
        
        # Sử dụng eval để tính toán kết quả tạm thời
        intermediate_result = eval(f"b {operations1} c {operations2} d")
        if intermediate_result > 0:
            break
        break
    # Tính toán biểu thức a*(b operations1 c operations2 d)
    result = eval(f"a * (1 + b {operations1} c {operations2} d)")

    # In ra biểu thức và kết quả dưới dạng LaTeX
    latex_a = fraction_to_latex(a)
    latex_expression = f"{latex_a} + {fraction_to_latex(b)} \\times {latex_a} {operations1}   {fraction_to_latex(c)} \\times {latex_a} {operations2} {fraction_to_latex(d)} \\times {latex_a}"
    latex_result = fraction_to_latex(result)
    return latex_expression, latex_result
# Tạo 4 phân số ngẫu nhiên khác nhau và không nguyên
fractions = [generate_non_integer_fraction() for _ in range(4)]
a, b, c, d = fractions
latex_expression, latex_result = calculate_and_print_latex_expression_2(a, b, c, d)  
doc.add_paragraph(f"$G= {latex_expression}$")
doc.add_paragraph(f"Đáp án: $G= {latex_result}$")
#

def calculate_and_print_latex_expression_3(a, b, c, d):
    # Tìm biểu thức (b operations1 c operations2 d) > 0
    while True:
        operations = [random.choice(['+', '-']) for _ in range(2)]
        operations1, operations2 = operations
        
        # Sử dụng eval để tính toán kết quả tạm thời
        intermediate_result = eval(f"b {operations1} c {operations2} d")
        if intermediate_result > 0:
            break
        break    
    # Tính toán biểu thức a*(b operations1 c operations2 d)
    result = eval(f" (b {operations1} c {operations2} d)/a")

    # In ra biểu thức và kết quả dưới dạng LaTeX
    latex_a = fraction_to_latex(a)
    latex_expression = f" {fraction_to_latex(b)} : {latex_a} {operations1}   {fraction_to_latex(c)} : {latex_a} {operations2} {fraction_to_latex(d)} : {latex_a}"
    latex_result = fraction_to_latex(result)
    return latex_expression, latex_result
# Tạo 4 phân số ngẫu nhiên khác nhau và không nguyên
fractions = [generate_non_integer_fraction() for _ in range(4)]
a, b, c, d = fractions
latex_expression, latex_result = calculate_and_print_latex_expression_3(a, b, c, d)  
doc.add_paragraph(f"$H= {latex_expression}$")
doc.add_paragraph(f"Đáp án: $H= {latex_result}$")
#
def calculate_and_print_latex_expression_4(a, b, c, d):
    # Tìm biểu thức (b operations1 c operations2 d) > 0
    while True:
        operations = [random.choice(['+', '-']) for _ in range(2)]
        operations1, operations2 = operations
        
        # Sử dụng eval để tính toán kết quả tạm thời
        intermediate_result = eval(f"b {operations1} c {operations2} d")
        if intermediate_result > 0:
            break
        break
    # Tính toán biểu thức a*(b operations1 c operations2 d)
    result = eval(f"( b/a {operations1} c/a {operations2} d/a)")

    # In ra biểu thức và kết quả dưới dạng LaTeX
    latex_a = fraction_to_latex(a)
    latex_expression = f"{latex_a} : {fraction_to_latex(b)} {operations1} {latex_a} :  {fraction_to_latex(c)} {operations2} {latex_a} : {fraction_to_latex(d)}"
    latex_result = fraction_to_latex(result)
    return latex_expression, latex_result
# Tạo 4 phân số ngẫu nhiên khác nhau và không nguyên
fractions = [generate_non_integer_fraction() for _ in range(4)]
a, b, c, d = fractions
latex_expression, latex_result = calculate_and_print_latex_expression_4(a, b, c, d)  
doc.add_paragraph(f"$K= {latex_expression}$")
doc.add_paragraph(f"Đáp án: $K= {latex_result}$")

#
###
import random
from fractions import Fraction
#
def calculate_L_1(n):
    k = random.randint(2, 5)
    L = Fraction(1, 1)  # Khởi tạo L là 1 để thực hiện phép nhân
    expression = ""  # Chuỗi biểu diễn biểu thức L
    for i in range(k, n + k):
        L *= Fraction(1, i) + 1
        if i == k:
            expression += f"(\\frac{{1}}{{{i}}} + 1)"
        else:
            expression += f" \\times (\\frac{{1}}{{{i}}} + 1)"
    return L, expression

# Giả sử n là một giá trị ngẫu nhiên từ 10 đến 15
n = random.randint(4, 6)
L, expression = calculate_L_1(n)

# In ra biểu thức và kết quả
doc.add_paragraph(f"$L = {expression}$")
if L.denominator == 1:
    doc.add_paragraph(f"Đáp án: $L= {int(L)}$")
else:
    doc.add_paragraph(f"Đáp án: $L= \\frac{{{L.numerator}}}{{{L.denominator}}}$")

n = random.randint(4, 7)
L, expression = calculate_L_1(n)

# In ra biểu thức và kết quả
doc.add_paragraph(f"$M = {expression}$")
if L.denominator == 1:
    doc.add_paragraph(f"Đáp án: $M = {int(L)}$")
else:
    doc.add_paragraph(f"Đáp án: $M = \\frac{{{L.numerator}}}{{{L.denominator}}}$")

n = random.randint(8, 12)
L, expression = calculate_L_1(n)

# In ra biểu thức và kết quả
doc.add_paragraph(f"$N = {expression}$")
if L.denominator == 1:
    doc.add_paragraph(f"Đáp án: $N = {int(L)}$")
else:
    doc.add_paragraph(f"Đáp án: $N = \\frac{{{L.numerator}}}{{{L.denominator}}}$")
#


#
import random
from fractions import Fraction

from fractions import Fraction

def calculate_L(n):
    L = Fraction(1, 1)  # Khởi tạo L là 1 để thực hiện phép nhân
    if n <= 4:
        expression = "".join(f" \\times (\\frac{{1}}{{{i}}} + 1)" if i != 2 else f"(\\frac{{1}}{{{i}}} + 1)" for i in range(2, n + 1))
    else:
        expression = "(\\frac{1}{2} + 1) \\times (\\frac{1}{3} + 1) \\times (\\frac{1}{4} + 1) + \\cdots + (\\frac{1}{" + str(n) + "} + 1)"
    for i in range(2, n + 1):
        L *= Fraction(1, i) + 1
    return L, expression

# Giả sử n là một giá trị ngẫu nhiên từ 30 đến 2025
n = random.randint(50, 100)
L, expression = calculate_L(n)

# In ra biểu thức và kết quả
doc.add_paragraph(f"$O = {expression}$")
if L.denominator == 1:
    doc.add_paragraph(f"Đáp án: $O= {int(L)}$")
else:
    doc.add_paragraph(f"Đáp án: $O= \\frac{{{L.numerator}}}{{{L.denominator}}}$")

n = random.randint(30, 2025)
L, expression = calculate_L(n)

# In ra biểu thức và kết quả
doc.add_paragraph(f"$P = {expression}$")
if L.denominator == 1:
    doc.add_paragraph(f"Đáp án: $P= {int(L)}$")
else:
    doc.add_paragraph(f"Đáp án: $P= \\frac{{{L.numerator}}}{{{L.denominator}}}$")

n = random.randint(2020, 2025)
L, expression = calculate_L(n)

# In ra biểu thức và kết quả
doc.add_paragraph(f"$Q = {expression}$")
if L.denominator == 1:
    doc.add_paragraph(f"Đáp án: $Q= {int(L)}$")
else:
    doc.add_paragraph(f"Đáp án: $Q= \\frac{{{L.numerator}}}{{{L.denominator}}}$")
#
#
#
#
# Bài 3
doc.add_heading("Tính nhanh.", level=3)
from fractions import Fraction
#
#
def calculate_A(n):
    k= random.randint(2, 5)
    A = Fraction(0)  # Khởi tạo A là 0
    expression = ""
    for i in range(k, n + k):  # Điều chỉnh để bao gồm n trong vòng lặp
        A += Fraction(1, i * (i + 1))
        if n <= 12 or i <= 3+k:
            if expression:  # Nếu expression không rỗng, thêm dấu '+'
                expression += " + "
            expression += f"\\frac{{1}}{{{i} \\times {i + 1}}}"
        elif i == 4+k:
            expression += " + \\cdots "
    if n > 12:
        expression += f" + \\frac{{1}}{{{n - 1} \\times {n}}}"

    # Định dạng kết quả
    if A.denominator == 1:  # Nếu A là số nguyên
        result = f"{int(A)}"
    else:  # Nếu A là phân số
        result = f"\\frac{{{A.numerator}}}{{{A.denominator}}}"

    return expression, result

# Ví dụ sử dụng
n = 7  # Giả sử n là 7
expression, result = calculate_A(n)
doc.add_paragraph(f"$A = {expression}$")
doc.add_paragraph(f"Đáp án: $A = {result}$")

# Ví dụ sử dụng
n = 9  # Giả sử n là 9
expression, result = calculate_A(n)
doc.add_paragraph(f"$B = {expression}$")
doc.add_paragraph(f"Đáp án: $B = {result}$")

n = 12  # Giả sử n là 12
expression, result = calculate_A(n)
doc.add_paragraph(f"$C = {expression}$")
doc.add_paragraph(f"Đáp án: $C = {result}$")

n = 15  # Giả sử n là 15
expression, result = calculate_A(n)
doc.add_paragraph(f"$D = {expression}$")
doc.add_paragraph(f"Đáp án: $D = {result}$")

n = 20  # Giả sử n là 20
expression, result = calculate_A(n)
doc.add_paragraph(f"$E = {expression}$")
doc.add_paragraph(f"Đáp án: $E = {result}$")

n = 99  # Giả sử n là 99
expression, result = calculate_A(n)
doc.add_paragraph(f"$F = {expression}$")
doc.add_paragraph(f"Đáp án: $F = {result}$")

#
#

def calculate_B(n):
    A = Fraction(0)  # Khởi tạo A là 0
    expression = ""
    k= random.randint(2, 5)
    for i in range(k, n + k):  # Điều chỉnh để bao gồm n trong vòng lặp
        A += Fraction(2, i * (i + 1))
        if n <= 12 or i <= 3+k:
            if expression:  # Nếu expression không rỗng, thêm dấu '+'
                expression += " + "
            expression += f"\\frac{{2}}{{{i} \\times {i + 1}}}"
        elif i == 4+k:
            expression += " + \\cdots "
    if n > 12:
        expression += f" + \\frac{{2}}{{{n - 1} \\times {n}}}"

    # Định dạng kết quả
    if A.denominator == 1:  # Nếu A là số nguyên
        result = f"{int(A)}"
    else:  # Nếu A là phân số
        result = f"\\frac{{{A.numerator}}}{{{A.denominator}}}"

    return expression, result

# Ví dụ sử dụng
n = 8  # Giả sử n là 8
expression, result = calculate_B(n)
doc.add_paragraph(f"$G = {expression}$")
doc.add_paragraph(f"Đáp án: $G = {result}$")

n = 10  # Giả sử n là 10
expression, result = calculate_B(n)
doc.add_paragraph(f"$H = {expression}$")
doc.add_paragraph(f"Đáp án: $H = {result}$")

n = 30  # Giả sử n là 30
expression, result = calculate_B(n)
doc.add_paragraph(f"$I = {expression}$")
doc.add_paragraph(f"Đáp án: $I = {result}$")

#

#

def calculate_C(n):
    k= random.randint(2, 5)
    A = Fraction(0)  # Khởi tạo A là 0
    expression = ""
    for i in range(k, 2*(n+k) + 1, 2):  # Điều chỉnh để bao gồm n trong vòng lặp
        A += Fraction(2, i * (i + 2))
        if n <= 7 or i <= 6+k:
            if expression:  # Nếu expression không rỗng, thêm dấu '+'
                expression += " + "
            expression += f"\\frac{{2}}{{{i} \\times {i + 2}}}"
        elif i == 8+k:
            expression += " + \\cdots "
    if n > 7:
        expression += f" + \\frac{{2}}{{{2*(n+k)+k} \\times {2*(n+k)+k+2}}}"

    # Định dạng kết quả
    if A.denominator == 1:  # Nếu A là số nguyên
        result = f"{int(A)}"
    else:  # Nếu A là phân số
        result = f"\\frac{{{A.numerator}}}{{{A.denominator}}}"

    return expression, result

# Ví dụ sử dụng
n = 5  # Giả sử n là 8
expression, result = calculate_C(n)
doc.add_paragraph(f"$J = {expression}$")
doc.add_paragraph(f"Đáp án: $J = {result}$")

n = 7  # Giả sử n là 7
expression, result = calculate_C(n)
doc.add_paragraph(f"$K = {expression}$")
doc.add_paragraph(f"Đáp án: $K = {result}$")

n = 10  # Giả sử n là 10
expression, result = calculate_C(n)
doc.add_paragraph(f"$L = {expression}$")
doc.add_paragraph(f"Đáp án: $L = {result}$")

n = 15  # Giả sử n là 15
expression, result = calculate_C(n)
doc.add_paragraph(f"$M = {expression}$")
doc.add_paragraph(f"Đáp án: $M = {result}$")

n = 30  # Giả sử n là 30
expression, result = calculate_C(n)
doc.add_paragraph(f"$N = {expression}$")
doc.add_paragraph(f"Đáp án: $N = {result}$")

n = 99  # Giả sử n là 99
expression, result = calculate_C(n)
doc.add_paragraph(f"$O = {expression}$")
doc.add_paragraph(f"Đáp án: $O = {result}$")

n = 2025  # Giả sử n là 2025
expression, result = calculate_C(n)
doc.add_paragraph(f"$P = {expression}$")
doc.add_paragraph(f"Đáp án: $P = {result}$")
#
#
#
#
#
def calculate_D(n):
    A = Fraction(0)  # Khởi tạo A là 0
    expression = ""
    k = random.randint(2, 5)
    h = random.randint(2, 5)
    for i in range(k, h*(n+k) + 1, h):  # Điều chỉnh để bao gồm n trong vòng lặp
        A += Fraction(h, i * (i + h))
        if n <= 5 or i <= 3*h+k:
            if expression:  # Nếu expression không rỗng, thêm dấu '+'
                expression += " + "
            expression += f"\\frac{{{h}}}{{{i} \\times {i + h}}}"
        elif i == 4*h+k:
            expression += " + \\cdots "
    if n > 5:
        expression += f" + \\frac{{{h}}}{{{h*(n+k)+k} \\times {h*(n+k)+h+k}}}"

    # Định dạng kết quả
    if A.denominator == 1:  # Nếu A là số nguyên
        result = f"{int(A)}"
    else:  # Nếu A là phân số
        result = f"\\frac{{{A.numerator}}}{{{A.denominator}}}"

    return expression, result

# Ví dụ sử dụng
n = 6  # Giả sử n là 6
expression, result = calculate_D(n)
doc.add_paragraph(f"$Q = {expression}$")
doc.add_paragraph(f"Đáp án: $Q = {result}$")

n = 8  # Giả sử n là 8
expression, result = calculate_D(n)
doc.add_paragraph(f"$R = {expression}$")
doc.add_paragraph(f"Đáp án: $R = {result}$")

n = 50  # Giả sử n là 50
expression, result = calculate_D(n)
doc.add_paragraph(f"$S = {expression}$")
doc.add_paragraph(f"Đáp án: $S = {result}$")

n = 197  # Giả sử n là 197
expression, result = calculate_D(n)
doc.add_paragraph(f"$T = {expression}$")
doc.add_paragraph(f"Đáp án: $T = {result}$")

n = 666  # Giả sử n là 2022
expression, result = calculate_D(n)
doc.add_paragraph(f"$U = {expression}$")
doc.add_paragraph(f"Đáp án: $U = {result}$")

n = 671  # Giả sử n là 2025
expression, result = calculate_D(n)
doc.add_paragraph(f"$V = {expression}$")
doc.add_paragraph(f"Đáp án: $V = {result}$")



#

#
def calculate_E(n):
    k = random.randint(2, 5)
    h = random.randint(2, 5)
    m = random.randint(4, 9)
    A = Fraction(0)  # Khởi tạo A là 0
    expression = ""
    for i in range(k, h*(n+k) + 1, h):  # Điều chỉnh để bao gồm n trong vòng lặp
        A += Fraction(m, i * (i + h))
        if n <= 5 or i <= 3*h+k:
            if expression:  # Nếu expression không rỗng, thêm dấu '+'
                expression += " + "
            expression += f"\\frac{{{m}}}{{{i} \\times {i + h}}}"
        elif i == 4*h+k:
            expression += " + \\cdots "
    if n > 5:
        expression += f" + \\frac{{{m}}}{{{h*(n+k)+k} \\times {h*(n+k)+k+4}}}"

    # Định dạng kết quả
    if A.denominator == 1:  # Nếu A là số nguyên
        result = f"{int(A)}"
    else:  # Nếu A là phân số
        result = f"\\frac{{{A.numerator}}}{{{A.denominator}}}"

    return expression, result

# Ví dụ sử dụng
n = 4  # Giả sử n là 6
expression, result = calculate_E(n)
doc.add_paragraph(f"$W = {expression}$")
doc.add_paragraph(f"Đáp án: $W = {result}$")

n = 20  # Giả sử n là 20
expression, result = calculate_E(n)
doc.add_paragraph(f"$X = {expression}$")
doc.add_paragraph(f"Đáp án: $X = {result}$")

n = 30  # Giả sử n là 30
expression, result = calculate_E(n)
doc.add_paragraph(f"$Y = {expression}$")
doc.add_paragraph(f"Đáp án: $Y = {result}$")

n = 502  # Giả sử n là 2022
expression, result = calculate_E(n)
doc.add_paragraph(f"$Z = {expression}$")
doc.add_paragraph(f"Đáp án: $Z = {result}$")

###########
#


#############################################################################
import os
import datetime
doc.add_paragraph("Phiếu bài tập được ra số liệu ngẫu nhiền bằng python + TeX")

# Get the current directory
current_directory = os.path.dirname(os.path.abspath(__file__))

# Get the current python file name
current_file_name = os.path.splitext(os.path.basename(__file__))[0]

# Format the current time as a string
current_time = datetime.datetime.now().strftime("%H%M%S_%d%m%Y")

# Save the document with the current time as the filename in the current directory
# Save the document
file_path = os.path.join(current_directory, f'{current_file_name}_{current_time}.docx')
doc.save(file_path)

# Open the document
os.startfile(file_path)
